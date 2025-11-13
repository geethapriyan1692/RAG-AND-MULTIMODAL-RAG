import json
from fastapi import APIRouter
from fastapi import Form, UploadFile
from typing import Any, Dict
from CenturyAssistant.app_config import config_loader
import logging
import traceback
from langchain_milvus import Milvus, BM25BuiltInFunction
from langchain_huggingface import HuggingFaceEmbeddings
from pymilvus import MilvusClient
import requests
from datetime import datetime
import cx_Oracle
import datetime as dt
from fastapi.responses import StreamingResponse
import re
import uuid
from langchain_ollama import OllamaLLM
from pymilvus import connections, CollectionSchema, FieldSchema, DataType, Collection
from PIL import Image
from io import BytesIO
import numpy as np
import base64
import fitz
from langchain_unstructured import UnstructuredLoader
from langchain_core.documents import Document
from transformers import CLIPProcessor, CLIPModel
import tempfile
import os
import win32com.client
from PIL import ImageGrab
from docx import Document as DocxDocument
from langchain.document_loaders import TextLoader

clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")


router = APIRouter()
LOGGER = logging.getLogger(__name__)

connection_str = config_loader.get("milvus_url")
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
static_collection_name = config_loader.get("collection_name")
vectorstore = Milvus(
    collection_name=static_collection_name,
    embedding_function=embeddings,
    builtin_function=BM25BuiltInFunction(),
    vector_field=["dense", "sparse"],
    connection_args={
        "uri": connection_str,
    },
)
llm = OllamaLLM(model="llama3.3:latest", base_url="http://prod-llm.excelacomcloud.net:11434")

db_user = config_loader.get("ml_username")
db_pass = config_loader.get("ml_password")
db_ip = config_loader.get("ml_catalog_host")
db_port = config_loader.get("ml_port")
db_service_name = config_loader.get("ml_service_name")

@router.post("/ragCollections/")
def rag_collections(request: Dict[Any, Any]):
    try:
        LOGGER.info("RAG Collections Method - Starts")
        userName = request['userName']
        LOGGER.info(userName)
        connections.connect(uri=connection_str)
 
        collection_name = config_loader.get("collection_name")
 
        namespaces = partition_retrieval(userName)
        LOGGER.info(f"Partition Retrieved : {namespaces}")
       
        response = {
            "collections": namespaces,
            "session_id": uuid.uuid4()
        }
 
        LOGGER.info("RAG Collections Method - Ends")
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())


@router.post("/ragFilterCollections/")
def search_collections(request: Dict[Any, Any]):
    try:
        """
        Author : Abhishek P
        Created On : 10/02/2025
        This method is for searching the collection in Milvus.
        """
        # Getting the searchValue from the request
        search_value = request["searchValue"]
        LOGGER.info("RAG filter Collections Method - Starts")
        # Connecting to Milvus DB
        
        client = MilvusClient(uri=connection_str)
        # Getting all the collections
        filtered_Partitions = []
 
        collection_name = config_loader.get("collection_name")
       
        results = client.query(
            collection_name=collection_name,
            filter="pk >= 0",  
            output_fields=["namespace"],  
            limit=1000  
        )
 
        namespaces = sorted(set(row["namespace"] for row in results))  
       
        for partition in namespaces:
            if partition.lower().startswith(search_value.lower()):
                filtered_Partitions.append(partition)
       
        response = {
            "collections": filtered_Partitions
        }
        LOGGER.info("RAG filter Collections Method - Ends")
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())

def save_chat_history(username,session_id,message_type,message,llm_method):
    try:
        """
        Author : Abhishek P
        Created On : 10/03/2025
        This method is to save chat history to the DB.
        """
        LOGGER.info("Save Chat History Method - Starts")

        # Connecting with the DB
        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
 
        message_time = datetime.now()
        # Cursor Creation
        cursor = connection.cursor()
        if message_type == "user":
            llm_endpoint = "http://prod-llm.excelacomcloud.net:11434/api/generate"
            payload = {
                "prompt":"Summarize the user query into a concise and meaningful chat title of 3 to 4 words maximum. Focus only on the core intent using natural, readable language. Avoid long or detailed phrases. If the user is simply greeting, summarize as 'Greeting Message'. Do not add any extra content or formatting.user_query- {}".format(message),
                "model":"llama3.2-vision:latest",
                "stream": False
            }
            response_value = requests.post(llm_endpoint, json=payload)
            response_data = json.loads(response_value.text)
            summarizer = response_data["response"]
            sql_insert = """INSERT INTO message_history (username, session_id, message_type, message, message_time, summarised_message, llm_method) VALUES (:1, :2, :3, :4, :5, :6, :7)"""
            # Values to the inserted into the DB
            data_to_insert = [(username, session_id, message_type, message, message_time, summarizer, llm_method)]
        else:
            # Insertion query
            sql_insert = """INSERT INTO message_history (username, session_id, message_type, message, message_time, llm_method) VALUES (:1, :2, :3, :4, :5, :6)"""
            # Values to the inserted into the DB
            data_to_insert = [(username, session_id, message_type, message, message_time, llm_method)]
        # Executing the query
        cursor.executemany(sql_insert, data_to_insert)
        # Commit the transaction
        connection.commit()
 
        response = {
            "status":"Success"
        }
 
        LOGGER.info("Save Chat History Method - Ends")
        return response
 
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    finally:
        # Closing the DB Connection
        connection.close()


@router.post("/getChatHistory/")
def get_chat_history(request: Dict[Any,Any]):
    try:
        """
        Author : Abhishek P
        Created On : 10/03/2025
        This method is to retrieve the chat history from the DB.
        """
        LOGGER.info("Get Chat History Method - Starts")

        # Connecting with the DB
        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)

        # Getting the username from the request
        username = request["username"]
        # Creating the Cursor
        cursor = connection.cursor()
        # SQL Query for retrieving the values based on the username  
        sql_select = """SELECT username, session_id, message_type, message, message_time FROM message_history 
        WHERE username = :1 ORDER BY message_time DESC """

        # Executing the Query
        cursor.execute(sql_select, (username,))
        # Retrieving all the values matching the condition
        fetched_values = cursor.fetchall()
        # Retrieving the column names
        column_names = [desc[0] for desc in cursor.description]
        
        # Pairing the fetched values with the column_names
        response= []
        for row in fetched_values:
            row_dict = {}
            row1 = [str(row_val) for row_val in row]
            for col, val in zip(column_names, row1):
                if col == "message":
                    row_dict[col] = str(val)  
                else:
                    row_dict[col] = val
            response.append(row_dict)

        LOGGER.info("Get Chat History Method - Ends")
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    finally:
        # Closing the DB Connection
        connection.close()


def get_chat_session_history(username,session_id):
    try:
        """
        Author : Abhishek P
        Created On : 10/03/2025
        This method is to retrieve the chat history from the DB.
        """
        LOGGER.info("Get Chat History Method - Starts")
 
        # Connecting with the DB
        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
 
        # Creating the Cursor
        cursor = connection.cursor()
        # SQL Query for retrieving the values based on the username  
        sql_select = """
        SELECT
            message_type,
            message
        FROM
            (
                SELECT
                    message_type,
                    message,
                    DENSE_RANK()
                    OVER(
                        ORDER BY
                            message_time DESC
                    ) AS dk
                FROM
                    message_history
                WHERE
                        username = :1
                    AND session_id = :2
            )
        WHERE
            dk < 7
        """
  
        # Executing the Query
        cursor.execute(sql_select, (username,session_id,))
        # Retrieving all the values matching the condition
        fetched_values = cursor.fetchall()
        # Retrieving the column names
        column_names = [desc[0] for desc in cursor.description]
       
        # Pairing the fetched values with the column_names
        response = [dict(zip(column_names, row)) for row in fetched_values]
        response_str = json.dumps(response, default=str)
 
        LOGGER.info("Get Chat History Method - Ends")
        return response_str
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    finally:
        # Closing the DB Connection
        connection.close()


@router.post("/retrieveChatHistory/")
def retrieve_chat_history(request: Dict[Any,Any]):
    try:
        session_id = request['sessionId']
        username = request['userName']

        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
        
        
        # Creating the Cursor
        cursor = connection.cursor()
        # SQL Query for retrieving the values based on the username  
        sql_select = """SELECT message_type, message, message_time, llm_method FROM message_history
        WHERE session_id = :1 ORDER BY message_time """
 
        # Executing the Query
        cursor.execute(sql_select, (session_id,))
        # Retrieving all the values matching the condition
        fetched_values = cursor.fetchall()
        # Retrieving the column names
        column_names = [desc[0] for desc in cursor.description]
       
        
        response= []
        for row in fetched_values:
            row_dict = {}
            row1 = [str(row_val) for row_val in row]
            for col, val in zip(column_names, row1):
                if col == "message":
                    row_dict[col] = str(val)  
                else:
                    row_dict[col] = val
            response.append(row_dict)
        
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    finally:
        # Closing the DB Connection
        connection.close()


@router.post("/showSessionHistory/")
def show_session_history(request: Dict[Any,Any]):
    try:
        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
        
        
        # SQL Query for retrieving the values based on the username  
        sql_select = """SELECT
            message_type,
            message,
            message_time,
            session_id,
            summarised_message
        FROM
            (
                SELECT
                    message_type,
                    message,
                    message_time,
                    session_id,
                    summarised_message,
                    DENSE_RANK() OVER(
                        PARTITION BY session_id
                        ORDER BY
                            message_time
                    ) AS dk
                FROM
                    message_history
                """
        
        cursor = connection.cursor()

        if 'llm_method' in request.keys() and 'username' in request.keys():
            llm_method = request['llm_method']
            username = request['username']
            
            if llm_method !="":
                sql_select = sql_select + """WHERE
                        message_type = 'user'
                        and llm_method = '{}'
                        and username = '{}'
                )
            WHERE
                dk = 1
            ORDER BY
                message_time DESC""".format(llm_method,username)
            else:
                sql_select = sql_select + """WHERE
                        message_type = 'user'
                        and username = '{}'
                )
            WHERE
                dk = 1
            ORDER BY
                message_time DESC""".format(username)
            cursor.execute(sql_select)
        else:
            if 'username' in request.keys():
                llm_method = 'rag'
                username = request['username']
                sql_select = """SELECT
                    message_type,
                    message,
                    message_time,
                    session_id,
                    summarised_message
                FROM
                    (
                        SELECT
                            message_type,
                            message,
                            message_time,
                            session_id,
                            summarised_message,
                            DENSE_RANK() OVER(
                                PARTITION BY session_id
                                ORDER BY
                                    message_time
                            ) AS dk
                        FROM
                            message_history
                        WHERE
                            message_type = 'user'
                            and llm_method = '{}'
                            and username = '{}'
                    )
                WHERE
                    dk = 1
                ORDER BY
                    message_time DESC""".format(llm_method, username)
                
            else:
                llm_method = 'rag'
                sql_select = """SELECT
                    message_type,
                    message,
                    message_time,
                    session_id,
                    summarised_message
                FROM
                    (
                        SELECT
                            message_type,
                            message,
                            message_time,
                            session_id,
                            summarised_message,
                            DENSE_RANK() OVER(
                                PARTITION BY session_id
                                ORDER BY
                                    message_time
                            ) AS dk
                        FROM
                            message_history
                        WHERE
                            message_type = 'user'
                            and llm_method = '{}'
                    )
                WHERE
                    dk = 1
                ORDER BY
                    message_time DESC""".format(llm_method)
                
            # Executing the Query
            cursor.execute(sql_select)
        
        # Retrieving all the values matching the condition
        fetched_values = cursor.fetchall()
        # Retrieving the column names
        column_names = [desc[0] for desc in cursor.description]

        history_data= []
        for row in fetched_values:
            row_dict = {}
            row1 = [str(row_val) for row_val in row]
            for col, val in zip(column_names, row1):
                if col == "message":
                    row_dict[col] = str(val)  
                else:
                    row_dict[col] = val
            history_data.append(row_dict)
        current_time = datetime.now()
        final_history_data = []        
        from collections import defaultdict
        grouped_messages = defaultdict(list)
        for message in history_data:
            message_time = datetime.fromisoformat(str(message["MESSAGE_TIME"]))
            
            # Calculate the difference in days
            difference = (current_time - message_time).days
            
            # Determine the value to add based on the difference
            if difference == 0:
                message["TIME_LABEL"] = "Today"
            elif difference == 1:
                message["TIME_LABEL"] = "Yesterday"
            elif difference == 2:
                message["TIME_LABEL"] = "2 days before"
            else:
                message["TIME_LABEL"] = "Older"

            final_history_data.append(message)
            time_label = message["TIME_LABEL"]
            grouped_messages[time_label].append({
                "MESSAGE_TYPE": message["MESSAGE_TYPE"],
                "MESSAGE": message["MESSAGE"],
                "MESSAGE_TIME": message["MESSAGE_TIME"],
                "SESSION_ID": message["SESSION_ID"],
                "SUMMARISED_MESSAGE": message["SUMMARISED_MESSAGE"]
            })
            
        desired_order = []
        if 'Today' in grouped_messages.keys():
            desired_order.append('Today')
        if 'Yesterday' in grouped_messages.keys():
            desired_order.append('Yesterday')
        if '2 days before' in grouped_messages.keys():
            desired_order.append('2 days before')
        if 'Older' in grouped_messages.keys():
            desired_order.append('Older')
        LOGGER.info(desired_order)
 
        ordered_dict = {key: grouped_messages[key] for key in desired_order}
        ordered_dict['session_id'] = uuid.uuid4()
        return ordered_dict
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    finally:
        # Closing the DB Connection
        connection.close()

@router.post("/removeSessionHistory/")
def remove_session_history(request: Dict[Any,Any]):
    try:
        session_id = request['sessionId']
        username = request['userName']

        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
        
        
        # Creating the Cursor
        cursor = connection.cursor()
        # SQL Query for retrieving the values based on the username  
        sql_select = """DELETE FROM message_history
        WHERE username = :1 and session_id = :2 """
 
        # Executing the Query
        cursor.execute(sql_select, (username,session_id,))
        connection.commit()
               
        # Pairing the fetched values with the column_names
        response = {'status': 'Success','message':'Session deleted successfully!'}
        
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    finally:
        # Closing the DB Connection
        connection.close()

@router.post("/generateSessionId/")
def generate_session_id():
    try:
        unique_id = uuid.uuid4()
        response={
            "status":"Success",
            "session_id": unique_id
        }
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())


def rephrase_user_query(user_query, convo_dict):
    try:
        """
        Author : Abhishek P
        Created On : 27/06/2025
        This method is to rephrase the user query.
        """
        LOGGER.info("Rephrasing User Query -  Starts")

        query_rephrase_prompt_starttime = datetime.now()
        convo_dict = json.loads(convo_dict)

        user_messages = [d["MESSAGE"] for d in convo_dict if d.get("MESSAGE_TYPE") == "user" and "MESSAGE" in d]

        LOGGER.info(f"{user_messages}")
        
        template = """

        You are a precise and efficient Query Generator. Your task is to generate a single, meaningful user query based on the provided context.

        Input:

        The last 3 conversations (if available).
        The current user query.
        Instructions:

        Dependency Check:
        If the current user query is dependent on or relevant to the last user query in the provided conversations, generate a new, concise, and meaningful user query that combines the context of the last user query with the current user query using key points ONLY.
        If the current user query is independent or irrelevant to the last user query, return the current user query exactly as provided without modification.
        Do not consider the entire last conversation, rephrase the current user query using key points only if it is dependent.
        Do not mention about the last conversation in the rephrased user query only consider the important key words from the last conversation.
        If there is no relevance between last conversation and current user query, return as it is without changing it.

        Output:
        Generate only one query.
        Ensure the query is clear, specific, and aligned with the user's intent.
        Do not add extraneous details, explanations, or assumptions beyond what is necessary to fulfill the query.
        Constraints:
        Avoid generating queries that are vague, overly broad, or misaligned with the provided context.
        Do not modify the current user query unless it is explicitly dependent on the last user query.
        If no previous conversations are provided, treat the current user query as independent and don't modify the user query.

        Task:
        Based on the provided input, generate a single user query following the above instructions.

        Important:
            RETURN ONLY THE USER QUERY AND NO EXPLANATIONS OR CONTEXT OR ALLOWED.
        Conversation: {conversation}
        Current User Query: {user_query}
        """.format(conversation = json.dumps(user_messages), user_query = user_query)
        query_rephrase_prompt_endtime = datetime.now()

        query_rephrase_llm_starttime = datetime.now()

        llm_res = llm(template)

        LOGGER.info(f"Rephrased Query: {llm_res}")

        query_rephrase_llm_endtime = datetime.now()
        LOGGER.info("Query rephrase LLM Time difference in seconds: {}".format(str((query_rephrase_llm_endtime - query_rephrase_llm_starttime).total_seconds())))
        LOGGER.info("Query rephrase Prompt Time difference in seconds: {}".format(str((query_rephrase_prompt_endtime - query_rephrase_prompt_starttime).total_seconds())))
    
        LOGGER.info("Rephrasing User Query -  Ends")
        return llm_res
    except Exception as e:
        LOGGER.error(traceback.format_exc())

def _delete_collections_logic(request:Dict[Any,Any]):
    try:
        """ Method to delete the collection as per user request"""
        # Getting the collection from the request
        delete_value = request["collectionName"]
        LOGGER.info("RAG Delete Collections Method - Starts")
        # Connecting to Milvus DB
        
        client = MilvusClient(uri=connection_str)
        collection_name = config_loader.get("collection_name")

        client.load_collection(collection_name=collection_name)
        results = client.query(collection_name=collection_name,
                            filter=f"namespace == '{delete_value}'",
                            output_fields=["namespace"])
        if results:
            client.delete(collection_name=collection_name,
                        filter=f"namespace == '{delete_value}'")

        response = {
            "response": f"Partition {delete_value} Deleted Successfully"
        }
        LOGGER.info("RAG Delete Collection Method - Ends")
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())

@router.post("/ragRemoveCollections/")
def delete_collections(request: Dict[Any, Any]):
    return _delete_collections_logic(request)


@router.post("/multimodalRagFileUpload/")
async def multimodalrag_upload_file(file: UploadFile,
                    sessionId: str = Form(...),
                    userName: str = Form(...),
                    collectionName: str = Form(...)):
    try:
        LOGGER.info("Partition Name Before Strip : {}file".format(collectionName))
        partitionName = collectionName.strip()
        LOGGER.info("Partition Name After Strip : {}file".format(partitionName))

        content = await file.read()
        yield_flag = True
        async def event_generator():
            try:
                LOGGER.info("RAG Upload File Method - Starts")

                #Sending the value to UI for progress bar
                yield f"data: 2%\n\n"
                connections.connect(uri=connection_str) 
                collection_name = config_loader.get("collection_name")

                client = MilvusClient(uri=connection_str)
                
                try:
                    results = client.query(
                        collection_name=collection_name,
                        filter="pk >= 0",  
                        output_fields=["namespace"],  
                    )
                except:
                    results = []
                    pass
 
                LOGGER.info(file.filename)
                # file_extension = file.filename.split('.')[-1]
                file_extension = os.path.splitext(file.filename)[1].lower()

                LOGGER.info(file_extension)

                yield f"data: 5%\n\n"
                
                #RegEx pattern for allowing only alphabets and underscore for partitionName
                LOGGER.info("RegEx Starts")
                pattern_regex = r'^[A-Za-z_ ]+$'
                pattern_check = re.match(pattern_regex, partitionName)
                LOGGER.info("RegEx Ends")

                if not pattern_check:
                    response = {
                        "message": "Please use only alphabets and underscores in the partition name",
                        "sessionId": sessionId,
                        "userName": userName,
                        "status":"Failure"
                    }
                    yield_flag = False
                    yield f"data: {json.dumps(response)}\n\n"
                    if not yield_flag:
                        return 

                if file_extension not in ['.docx','.pdf','.txt', '.doc']:
                    response = {
                        "message": "Unsupported file type. Please upload a .pdf/.docx/.txt/.doc file",
                        "sessionId": sessionId,
                        "userName": userName,
                        "status": "Failure"
                    }
                    #Sending the response to UI, if any error occurs.
                    yield_flag = False
                    yield f"data: {json.dumps(response)}\n\n"
                    if not yield_flag:
                        return 
        
                namespaces = set(row["namespace"] for row in results)
                LOGGER.info(namespaces)

                if partitionName == '':
                    response = {
                        "message": "Please provide a name",
                        "sessionId": sessionId,
                        "userName": userName,
                        "status":"Failure"
                    }
                    yield_flag = False
                    yield f"data: {json.dumps(response)}\n\n"
                    if not yield_flag:
                        return 

                # Handling partitioning scenario
                if partitionName in namespaces:
                    request_data = {"collectionName": partitionName}
                    result = _delete_collections_logic(request_data)
                    
                collectionName = config_loader.get("collection_name")

                yield f"data: 10%\n\n"
                if file_extension == '.txt':
                    all_data = image_extraction_from_txt(file)

                elif file_extension == '.docx':
                    all_data = image_extraction_from_docx(partitionName, content, file, file_extension)
                elif file_extension == '.doc':
                    all_data = image_extraction_from_doc(partitionName, content, file_extension)
                elif file_extension == '.pdf':
                    all_data = image_extraction_from_pdf(partitionName, content, file) 

                response = requests.post(
                    "http://prod-llm.excelacomcloud.net:8006/rag/api/multimodalRagImage/",         
                    data={ "partitionName": partitionName,"userName": userName, "fileContent": all_data}
                )

                response_json = response.json() 
                LOGGER.info(response_json)          
                docs = response_json["doc_obj"] 

                # Convert list of dicts into list of Document objects
                documents = [Document(page_content=d["page_content"],metadata={**d["metadata"],"source":collectionName, "userName": userName}) for d in docs]
                yield f"data: 70%\n\n"
                
                vectorstore = Milvus.from_documents(
                    collection_name=collectionName,
                    documents=documents,
                    embedding=embeddings,
                    builtin_function=BM25BuiltInFunction(),
                    vector_field=["dense", "sparse"],
                    connection_args={
                        "uri": connection_str,
                    },
                    consistency_level="Strong",
                    partition_key_field="namespace"
                )
 
                _ = map_partition_user(partitionName,userName)
             
                yield f"data: 100%\n\n"

                response = {
                    "message": "File uploaded successfully",
                    "sessionId": sessionId,
                    "userName": userName,
                    "status":"Success"
                }
                yield f"data: {json.dumps(response)}\n\n"

                LOGGER.info("RAG Upload File Method - Ends")
            except Exception as e:
                LOGGER.error(traceback.format_exc())
                response = {
                    "message": "Upload failed. Please try again",
                    "sessionId": sessionId,
                    "userName": userName,
                    "status":"Failure"
                }
                yield f"data: {json.dumps(response)}\n\n"


        return StreamingResponse(event_generator(), media_type="text/event-stream")
    
    except Exception as e:
        LOGGER.error(traceback.format_exc())
        response = {
            "message": "Upload failed. Please try again",
            "sessionId": sessionId,
            "userName": userName,
            "status":"Failure"
        }
        return response
    finally:
        await file.close()


def image_extraction_from_pdf(partitionName, content, file):
    try:
        LOGGER.info("Image Extraction from Documents Method - Starts")

        IMG_OUTPUT_DIR = os.path.join(config_loader.get("img_path"), partitionName)
        os.makedirs(IMG_OUTPUT_DIR, exist_ok=True)
        image_embeddings = []
        extracted_content_chunks = []
        global_image_count = 0
        file_extension = os.path.splitext(file.filename)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            temp_file.write(content)
            file_path = temp_file.name
 
        if file_extension == '.pdf':
            loader = fitz.open(file_path)

        for page_num in range(len(loader)):
            page = loader[page_num]
            text_blocks = page.get_text("blocks")
            images_on_page = []
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                try:
                    image_data = loader.extract_image(xref)
                    img_bytes = image_data["image"]
                    img_ext = image_data["ext"]
                    img_width = image_data["width"]
                    img_height = image_data["height"]
                    img_area = img_width * img_height
                    img_filename = f"{partitionName}_image_{global_image_count}.{img_ext}"
                    img_path = os.path.join(IMG_OUTPUT_DIR, img_filename)
                    LOGGER.info(img_path)
                    with open(img_path, "wb") as img_file:
                        img_file.write(img_bytes)
                    tag_type = img_filename
                   
                    # Convert image to embedding
                    img = Image.open(BytesIO(img_bytes)).convert("RGB")
                    inputs = clip_processor(images=img, return_tensors="pt")
                    image_embedding = clip_model.get_image_features(**inputs).detach().numpy()
                    image_embedding = image_embedding / np.linalg.norm(image_embedding, axis=1, keepdims=True)  # Normalize
                    image_embeddings.append({
                        "embedding": image_embedding[0].tolist(),
                        "filename": img_filename,
                        "metadata": {"Collection_Name": partitionName, "image_id": img_filename}
                    })
 
                    img_rects = page.get_image_rects(xref)
                    y_pos, x_pos = (img_rects[0].y0, img_rects[0].x0) if img_rects else (0, 0)
                    normalized_path = os.path.normpath(img_path).replace("\\", "/")
                    images_on_page.append((y_pos, x_pos, 0, f"[{normalized_path}]"))
                    global_image_count += 1
                except Exception:
                    continue
            elements = [(y0, x0, 0, text.strip()) for x0, y0, x1, y1, text, block_no, block_type in text_blocks if block_type == 0 and text.strip()]
            elements.extend(images_on_page)
            elements.sort(key=lambda e: (e[0], e[1]))
            page_content = " ".join(e[3] for e in elements) 
            extracted_content_chunks.append(f"PAGE_START_{page_num + 1} {page_content} PAGE_END_{page_num + 1}")
 
        LOGGER.info("Image Extraction from Documents Method - Ends")
        loader.close()
        LOGGER.info("".join(extracted_content_chunks))
        return "\n".join(extracted_content_chunks), 
       
    except Exception as e:
        LOGGER.error(traceback.format_exc())
        return "", []
    

def image_extraction_from_docx(partitionName, content, file, file_extension):
    image_embeddings = []
    image_counter = 1  
    IMG_OUTPUT_DIR = os.path.join(config_loader.get("img_path"), partitionName)
    os.makedirs(IMG_OUTPUT_DIR, exist_ok=True)

    # Save uploaded file to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(content)
        file_path = temp_file.name

    # Create output image directory if not exists

    doc = DocxDocument(file_path)
    output_parts = []

    
    image_map = {}
    for rel in doc.part._rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            image_name = rel.target_ref.split("/")[-1]
            image_ext = image_name.split(".")[-1]
            image_map[rel.rId] = (image_bytes, image_ext)

    # Process paragraphs and extract images with placeholders (with extension)
    for para in doc.paragraphs:
        run_content = ""
        for run in para.runs:
            drawing_elements = run._element.xpath('.//w:drawing')
            if drawing_elements:
                # Find image in rels by traversing the XML tree
                blip = run._element.xpath('.//a:blip')
                if blip and blip[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"):
                    rel_id = blip[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if rel_id in image_map:
                        image_bytes, image_ext = image_map[rel_id]
                        placeholder_name = f"{partitionName}_{image_counter}.{image_ext}"
                        image_path = os.path.join(IMG_OUTPUT_DIR, placeholder_name)
                        normalized_path = os.path.normpath(image_path).replace("\\", "/")
                        placeholder = f"[{normalized_path}]"
                        

                        # Save image
                        with open(image_path, "wb") as f:
                            f.write(image_bytes)

                        # Add to image_embeddings
                        image_embeddings.append({
                            "image_name": placeholder_name,
                            "image_bytes": image_bytes
                        })

                        run_content += placeholder
                        image_counter += 1
            else:
                run_content += run.text

        if run_content.strip():
            output_parts.append(run_content.strip())

    final_text = "\n\n".join(output_parts)
    return final_text


def image_extraction_from_doc(partitionName, content, file_extension):


    IMG_OUTPUT_DIR = os.path.join(config_loader.get("img_path"), partitionName)
    os.makedirs(IMG_OUTPUT_DIR, exist_ok=True)
    # Save uploaded file to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(content)
        file_path = temp_file.name

    word = win32com.client.Dispatch("Word.Application")
    word.DisplayAlerts = 0
    doc = word.Documents.Open(file_path)

    image_counter = 1
    output_text = ""
    image_paths = []  # Store paths of extracted images

    for para in doc.Paragraphs:
        para_text = para.Range.Text.strip()
        shapes = para.Range.InlineShapes

        if shapes.Count > 0:
            for i in range(1, shapes.Count + 1):
                shape = shapes.Item(i)
                img_filename = f"{partitionName}_image_{image_counter}.png"
                image_path = os.path.join(IMG_OUTPUT_DIR, img_filename)

                try:
                    shape.Range.CopyAsPicture()
                    img = ImageGrab.grabclipboard()
                    if img:
                        img.save(image_path, 'PNG')
                        image_paths.append(image_path)  
                        para_text += f" [{image_path}]"
                    else:
                        para_text += f" [IMAGE_FAILED_{image_counter}]"
                except Exception as e:
                    para_text += f" [IMAGE_FAILED_{image_counter}]"

                image_counter += 1

        output_text += para_text + "\n"

    doc.Close(False)
    word.Quit()
    return output_text 

@router.post("/multimodalRagChat/")
def multimodalrag_with_memory(request: Dict[Any,Any]):
    try:
        LOGGER.info("Rag Chat method - Starts")
        user_query = request['userQuery']
        session_id = request['sessionId']
        username = request['userName']
        hyper_params = request["hyperParameters"]["temperature"]


        if hyper_params < 1:        
            namespace_list = partition_retrieval(username)
            LOGGER.info(namespace_list)
 
            # Create filter expression
            namespace_list_str = ', '.join(f'"{ns}"' for ns in namespace_list)
            filter_expression = f"namespace IN [{namespace_list_str}]"

 
            convo_dict = get_chat_session_history(username,session_id)
            rephrased_query = rephrase_user_query(user_query, convo_dict)
 
            retriever_starttime = datetime.now()
            retrieved_chunks = vectorstore.similarity_search(
                query=rephrased_query,
                k=4,
                expr=filter_expression,
            )
            retriever_endtime = datetime.now()


            retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
            retrieved_chunks = retriever.invoke(
                rephrased_query,
            )
            LOGGER.info("Retrieved Chunks: {}".format(retrieved_chunks))
            

            retrieved_docs = []
            for retrieved_chunk in retrieved_chunks:
                retrieved_docs.append(retrieved_chunk.page_content)
            LOGGER.info(retrieved_docs)
      

            prompt_starttime = datetime.now()
            prompt_value = "You are a helpful assistant that prioritizes answering questions naturally. Follow these rules:\n\nQuery Type Detection:\nIf the <question> is a general knowledge query (e.g., science, history, greetings), ignore the <context> entirely and answer directly using your inherent knowledge.\nIf the <question> is specific (e.g., about documents, processes, or data), use the <context> only if it is relevant.\n\nResponse Workflow:\nGeneral Questions (e.g., \"How many planets are there?\", \"Hi!\"):\nAnswer conversationally without mentioning the <context>.\nDo not reference irrelevance to context (e.g., avoid phrases like \"This isn’t related to the context\").\n\nContext-Dependent Questions:\nIf the <context> contains sufficient information: Respond naturally without tampering the data, also do not change the language and wording of answer, if answer contains steps or process then completely give full answer, integrating details without citations or references.\nIf the <context> is irrelevant/insufficient: State, \"I don’t have enough information to address that.\"\nAmbiguous Questions: Politely ask for clarification (e.g., \"Could you specify which aspect you’re interested in?\").\n\nProhibited Actions:\nNever mention the <context>’s content (e.g., \"The context discusses software...\") unless explicitly required to justify a limitation.\nAvoid disclaimers like \"This isn’t in the context\" for general questions.\nAvoid giving unnecessary or extra things in response, only give\nDo not infer context from <Message_History> unless the user explicitly ties the current question to it.\n\nMessage History Format:\nThe message history will include up to the last three interactions in this JSON array format:\n[\n  {\"MESSAGE_TYPE\": \"user\", \"MESSAGE\": \"last user query\"},\n  {\"MESSAGE_TYPE\": \"llm\", \"MESSAGE\": \"last query generated response by LLM\"},\n  {\"MESSAGE_TYPE\": \"user\", \"MESSAGE\": \"last 2nd user query\"},\n  {\"MESSAGE_TYPE\": \"llm\", \"MESSAGE\": \"last 2nd query generated response by LLM\"},\n  {\"MESSAGE_TYPE\": \"user\", \"MESSAGE\": \"last 3rd user query\"},\n  {\"MESSAGE_TYPE\": \"llm\", \"MESSAGE\": \"last 3rd query generated response by LLM\"}\n]\n\nInput Structure:\nEvery query will be provided in three parts:\n<question> {question} </question> : User Query\n<Message_History> {history} </Message_History>: Last 3 chat History\n<context> {context} </context>: Retrieved RAG Chunks. Please do not return the data other than the response you are generating. Never return tags mentioned in the template. ALWAYS return the only the data. Do not mention about the chunks retrieved or about context or the question asked. Always include correct image placeholder[imagename] in response"

            multimodal_prompt_value = """The context contains text chunks with image placeholders like [filename].
                    - Do NOT modify, rename, or remove these placeholders if the context is from a document with images (e.g., .pdf, .docx, .doc).
                    - If the context is from a .txt file, ignore and exclude all image placeholders in the response.
                    - Preserve the exact positions of placeholders in the text for non-text files as they appear in the context.
                    - Ignore any images containing the word 'excelacom' or the 'excelacom' logo.
                    Additional Rules:
                    - Use the context to ground your response, citing relevant text and placeholders.
                    - If image embeddings are relevant, note their presence via placeholders but do not describe image content directly.
                    - Always bind the image placeholder[imagename.png|.jpg] to its corresponding keyword and position—do not separate or reorder them under any condition.
                    Rules for General Questions (e.g., greetings, science, history):
                    - Do NOT include any image placeholders in the response.
                    - Respond naturally without referencing or including placeholders like [filename].
                    - Do not mention the context or image placeholders in your response.
                    - If the question is general, answer directly without using the context or image placeholders.
                    - If the question is specific, use the context to provide a detailed answer, including relevant image placeholders.
            
                    """
            
            
            template = """
                {prompt_value}
                {multimodal_prompt_value}
                <question> {question} </question>
                <Message_History> {history} </Message_History>
                <context> {context} </context>""".format(prompt_value=prompt_value, multimodal_prompt_value=multimodal_prompt_value, context=json.dumps(retrieved_docs),question=rephrased_query, history=json.dumps(convo_dict))
            
            prompt_endtime = datetime.now()

            llm_starttime = datetime.now()
            llm_res = llm(template)
            llm_endtime = datetime.now()

            llm_res = llm_res.replace("\\\\","/")

            status_val = save_chat_history(username,session_id,"user",user_query,'rag')
            status_val = save_chat_history(username,session_id,"llm",llm_res,'rag')
            LOGGER.info(status_val)
        else:
            retrieved_docs = []

            llm_starttime = datetime.now()
            llm_res = llm(template)
            llm_endtime = datetime.now()

            llm_res = llm_res.replace("\\\\","/")

            LOGGER.info(llm_res)
            status_val = save_chat_history(username,session_id,"user",user_query,'rag')
            status_val = save_chat_history(username,session_id,"llm",llm_res,'rag')
            LOGGER.info(status_val)


        response = {
            "llmResponse": llm_res,
            "chunks": retrieved_docs,
            "sessionId": session_id,
            "userName": username,
            "llmMethod" : "rag"
        }

        LOGGER.info("Retriever Time difference in seconds: {}".format(str((retriever_endtime - retriever_starttime).total_seconds())))
        LOGGER.info("Prompt Time difference in seconds: {}".format(str((prompt_endtime - prompt_starttime).total_seconds())))
        LOGGER.info("LLM Call Time difference in seconds: {}".format(str((llm_endtime - llm_starttime).total_seconds())))
        LOGGER.info("Rag Chat method - Ends")
        return response
    except Exception as e:
        LOGGER.error(traceback.format_exc())


def partition_retrieval(userName):
    try:
        LOGGER.info("Partition Retrieval from Oracle DB Method - Starts")
        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
        cursor = connection.cursor()
        # SQL Query for retrieving the values based on the username  
        sql_select = """SELECT partition_name FROM partition_user_map_multimodal WHERE user_name = :1"""
        # Executing the Query
        cursor.execute(sql_select, (userName,))
        # Retrieving all the values matching the condition
        fetched_values = cursor.fetchall()
        # Retrieving the column names
        column_values = [value for row in fetched_values for value in row]
        namespaces = sorted(set(column_values))
 
        LOGGER.info("Partition Retrieval from Oracle DB Method - Ends")
        return namespaces
   
    except:
        namespaces = []
        LOGGER.error(traceback.format_exc())
        return namespaces
    finally:
        connection.close()

 
def map_partition_user(partition_name, username):
    try:
 
        LOGGER.info("Partition User Mapping Method - Starts")
        dsn = cx_Oracle.makedsn(db_ip, db_port, service_name=db_service_name)
        connection = cx_Oracle.connect(user=db_user, password=db_pass, dsn=dsn)
        cursor = connection.cursor()
        message_time = datetime.now()
        sql_insert = """INSERT INTO partition_user_map_multimodal (user_name, partition_name, created_date) VALUES (:1, :2, :3)"""
        data_to_insert = [(username, partition_name, message_time)]
 
        cursor.executemany(sql_insert, data_to_insert)
        connection.commit()
 
        LOGGER.info("Partition User Mapping Method - Ends")
 
    except Exception as e:
        LOGGER.error(traceback.format_exc())
    return True
 
from langchain.document_loaders import TextLoader

def image_extraction_from_txt(file):
    try:
        LOGGER.info("Image Extraction from Text Method - Starts")
        loader = UnstructuredLoader(file=file.file, mode="elements", strategy="fast", metadata_filename=file.filename)

        loader_doc = loader.load()
        all_data = ""
        for each_doc in loader_doc:
            all_data += each_doc.page_content

        LOGGER.info("Image Extraction from Text Method - Ends")
        
        LOGGER.info(f"Extracted text content: {all_data}")
        return all_data
    except Exception as e:
        LOGGER.error(traceback.format_exc())
        return ""  
